import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import numpy as np

# Folder to save graphs
graph_dir = "graphs"
os.makedirs(graph_dir, exist_ok=True)

# Word report path
report_path = "Performance_Report.docx"


# ==== Dummy Data for Graphs (replace with real metrics if available) ====
def generate_dummy_data():
    return {
        "Handwriting Complexity": ["Simple", "Moderate", "Complex"],
        "Accuracy": [92.5, 88.2, 80.7],
        "Text Length": [100, 300, 600],
        "WER": [5.1, 6.8, 9.3],
        "Image Resolution": ["300 DPI", "600 DPI", "1200 DPI"],
        "Processing Time": [450, 520, 640],
        "Configurations": ["Without PSO", "With PSO"],
        "OCR Accuracy": [85.0, 93.2],
        "Confidence Scores": np.random.normal(loc=0.8, scale=0.1, size=100),
        "Handwriting Type": ["Cursive", "Print", "Mixed"],
        "Latency": [900, 750, 820],
        "Page Sizes": ["A4", "A3", "Letter"],
        "FPS": [5.2, 3.8, 4.5]
    }


# ==== Plot & Save Functions ====
def save_plot(fig, filename):
    filepath = os.path.join(graph_dir, filename)
    fig.savefig(filepath, bbox_inches='tight')
    plt.close(fig)
    return filepath


def plot_graphs(data):
    paths = []

    # 1. Accuracy vs. Handwriting Complexity
    fig, ax = plt.subplots()
    ax.bar(data["Handwriting Complexity"], data["Accuracy"], color="skyblue")
    ax.set_title("Accuracy vs. Handwriting Complexity")
    ax.set_ylabel("Accuracy (%)")
    paths.append(save_plot(fig, "accuracy_vs_complexity.png"))

    # 2. WER vs. Text Length
    fig, ax = plt.subplots()
    ax.plot(data["Text Length"], data["WER"], marker='o', color="salmon")
    ax.set_title("Word Error Rate vs. Text Length")
    ax.set_xlabel("Text Length (words)")
    ax.set_ylabel("WER (%)")
    paths.append(save_plot(fig, "wer_vs_textlength.png"))

    # 3. Processing Time vs. Image Resolution
    fig, ax = plt.subplots()
    ax.bar(data["Image Resolution"], data["Processing Time"], color="purple")
    ax.set_title("Processing Time vs. Image Resolution")
    ax.set_ylabel("Time (ms)")
    paths.append(save_plot(fig, "time_vs_resolution.png"))

    # 4. OCR Accuracy Before and After PSO
    fig, ax = plt.subplots()
    ax.bar(data["Configurations"], data["OCR Accuracy"], color=["gray", "green"])
    ax.set_title("OCR Accuracy (Before vs. After PSO)")
    ax.set_ylabel("Accuracy (%)")
    paths.append(save_plot(fig, "accuracy_pso.png"))

    # 5. OCR Confidence Score Distribution
    fig, ax = plt.subplots()
    ax.hist(data["Confidence Scores"], bins=20, color="orange", edgecolor='black')
    ax.set_title("OCR Confidence Score Distribution")
    ax.set_xlabel("Confidence Score")
    ax.set_ylabel("Frequency")
    paths.append(save_plot(fig, "confidence_distribution.png"))

    # 6. Latency vs. Handwriting Type
    fig, ax = plt.subplots()
    ax.bar(data["Handwriting Type"], data["Latency"], color="teal")
    ax.set_title("Digitization Latency vs. Handwriting Style")
    ax.set_ylabel("Latency (ms)")
    paths.append(save_plot(fig, "latency_vs_style.png"))

    # 7. FPS vs. Page Size
    fig, ax = plt.subplots()
    ax.bar(data["Page Sizes"], data["FPS"], color="gold")
    ax.set_title("Frames Per Second vs. Document Page Size")
    ax.set_ylabel("FPS")
    paths.append(save_plot(fig, "fps_vs_pagesize.png"))

    return paths


# ==== Add Graphs to Word ====
def generate_word_report(graph_paths):
    doc = Document()
    doc.add_heading("Performance Evaluation Report", 0)

    for path in graph_paths:
        title = os.path.splitext(os.path.basename(path))[0].replace("_", " ").title()
        doc.add_heading(title, level=2)
        doc.add_picture(path, width=Inches(5.5))

    doc.save(report_path)
    print(f"\n✅ Report saved as: {report_path}")


# ==== CLI Menu ====
def cli_menu():
    while True:
        print("\n📊 PERFORMANCE REPORT MENU")
        print("1. Generate & Save Graphs")
        print("2. Create Word Report with Graphs")
        print("3. Exit")

        choice = input("Enter your choice: ").strip()

        if choice == "1":
            print("\nGenerating graphs...")
            data = generate_dummy_data()
            graph_paths = plot_graphs(data)
            print(f"{len(graph_paths)} graphs saved to '{graph_dir}/'")

        elif choice == "2":
            if not os.listdir(graph_dir):
                print("⚠️ No graphs found. Please generate them first.")
            else:
                graph_paths = [os.path.join(graph_dir, f) for f in os.listdir(graph_dir) if f.endswith(".png")]
                generate_word_report(graph_paths)

        elif choice == "3":
            print("Goodbye!")
            break

        else:
            print("Invalid choice. Please try again.")


if __name__ == "__main__":
    cli_menu()
